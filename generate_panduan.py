import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE

BASE_URL = "https://sdkh3-tartil.helmark-apps.com"
OUTPUT_DIR = "panduan"

def add_heading_custom(doc, text, level=1, color=None):
    """Add heading with consistent styling."""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    run.font.name = "Plus Jakarta Sans"
    run.font.size = Pt({1: 18, 2: 14, 3: 12}.get(level, 12))
    run.font.bold = True
    if color:
        run.font.color.rgb = RGBColor(*color)
    else:
        run.font.color.rgb = RGBColor(28, 25, 23)  # #1c1917
    heading.paragraph_format.space_before = Pt(12 if level == 1 else 10)
    heading.paragraph_format.space_after = Pt(6)
    return heading


def add_paragraph_custom(doc, text, bold=False, italic=False, size=11, color=(68, 64, 60), bullet=False):
    """Add paragraph with consistent styling."""
    p = doc.add_paragraph()
    if bullet:
        p.style = "List Bullet"
    run = p.add_run(text)
    run.font.name = "Inter"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(*color)
    run.font.bold = bold
    run.font.italic = italic
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    return p


def add_url(doc, url, label=None):
    """Add URL paragraph."""
    p = doc.add_paragraph()
    p.add_run("URL: ").font.name = "Inter"
    run = p.add_run(url if label is None else label)
    run.font.name = "Inter"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(12, 138, 95)
    run.font.underline = True
    p.paragraph_format.space_after = Pt(6)
    return p


def add_screenshot_placeholder(doc, description):
    """Add screenshot placeholder."""
    p = doc.add_paragraph()
    p.add_run("[Screenshot: ").font.name = "Inter"
    run = p.add_run(description)
    run.font.name = "Inter"
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(120, 113, 108)
    p.add_run("]").font.name = "Inter"
    p.paragraph_format.space_after = Pt(8)
    return p


def add_info_box(doc, title, lines, box_color=(240, 253, 244)):
    """Add an info box with title and bullet lines."""
    add_paragraph_custom(doc, title, bold=True, size=11, color=(6, 95, 67))
    for line in lines:
        add_paragraph_custom(doc, line, bullet=True, size=10, color=(68, 64, 60))


def add_steps(doc, steps):
    """Add numbered steps."""
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(f"{step}")
        run.font.name = "Inter"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(68, 64, 60)
        p.paragraph_format.space_after = Pt(4)


def create_cover(doc, title, subtitle):
    """Create title page."""
    # Spacer
    for _ in range(6):
        doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("TARTIL Pro")
    run.font.name = "Plus Jakarta Sans"
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.color.rgb = RGBColor(12, 138, 95)

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_para.add_run("Sistem Penilaian Tartil Online")
    run.font.name = "Inter"
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(120, 113, 108)
    subtitle_para.paragraph_format.space_after = Pt(30)

    doc_para = doc.add_paragraph()
    doc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = doc_para.add_run(title)
    run.font.name = "Plus Jakarta Sans"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(28, 25, 23)

    doc_para2 = doc.add_paragraph()
    doc_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = doc_para2.add_run(subtitle)
    run.font.name = "Inter"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(120, 113, 108)

    for _ in range(8):
        doc.add_paragraph()

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("SD Khadijah 3")
    run.font.name = "Inter"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(120, 113, 108)

    doc.add_page_break()


def create_toc(doc, sections):
    """Create a simple table of contents."""
    add_heading_custom(doc, "Daftar Isi", level=1)
    for i, section in enumerate(sections, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. {section}").font.name = "Inter"
        p.runs[0].font.size = Pt(11)
        p.runs[0].font.color.rgb = RGBColor(68, 64, 60)
        p.paragraph_format.space_after = Pt(4)
    doc.add_page_break()


def generate_admin_doc():
    doc = Document()

    # Set default font styles
    style = doc.styles["Normal"]
    style.font.name = "Inter"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(68, 64, 60)

    create_cover(doc, "Buku Panduan Admin", "Panduan Lengkap Penggunaan Sistem untuk Administrator")

    sections = [
        "Login dan Akses Awal",
        "Dashboard Admin",
        "Manajemen Tahun Ajaran & Semester",
        "Manajemen Data Siswa",
        "Manajemen Data Guru",
        "Manajemen Kelas Reguler",
        "Manajemen Kelas Tartil",
        "Jurnal, Absensi, dan Monitoring",
        "Tahfidz & Hafalan",
        "Ujian Munaqosyah",
        "Penilaian Rapor Internal",
        "Perpindahan & Kenaikan Kelas",
        "Pendampingan Orang Tua",
        "Audit, Statistik & Laporan",
        "System Setup & Maintenance",
    ]
    create_toc(doc, sections)

    # 1. Login
    add_heading_custom(doc, "1. Login dan Akses Awal", level=1)
    add_paragraph_custom(doc, "Admin mengakses aplikasi melalui halaman landing utama.")
    add_url(doc, f"{BASE_URL}/")
    add_steps(doc, [
        "Buka URL aplikasi.",
        "Masukkan email admin dan password.",
        "Klik tombol Masuk.",
        "Sistem akan mengarahkan ke dashboard admin.",
    ])
    add_screenshot_placeholder(doc, "Halaman login aplikasi TARTIL")

    # 2. Dashboard
    add_heading_custom(doc, "2. Dashboard Admin", level=1)
    add_url(doc, f"{BASE_URL}/admin/dashboard")
    add_paragraph_custom(doc, "Dashboard menampilkan ringkasan data utama: jumlah siswa, guru, kelas, semester aktif, notifikasi R2 cache yang perlu di-precalculate, dan akses cepat ke menu penting.")
    add_screenshot_placeholder(doc, "Dashboard admin")

    # 3. Tahun Ajaran & Semester
    add_heading_custom(doc, "3. Manajemen Tahun Ajaran & Semester", level=1)
    add_heading_custom(doc, "3.1 Tahun Ajaran", level=2)
    add_url(doc, f"{BASE_URL}/admin/tahun-ajaran")
    add_paragraph_custom(doc, "Halaman untuk membuat, menutup, dan melihat riwayat tahun ajaran.")
    add_steps(doc, [
        "Klik menu 'Tahun Ajaran'.",
        "Klik tombol 'Tambah Tahun Ajaran'.",
        "Isi nama tahun ajaran (misal: 2026/2027).",
        "Sistem otomatis membuat semester Ganjil dan Genap, menutup tahun ajaran lama, dan menaikkan kelas reguler siswa.",
    ])
    add_info_box(doc, "Aturan Penting:", [
        "Semua kelas reguler jenjang 1–5 yang berpenghuni wajib memiliki kelas tujuan sebelum tahun ajaran baru dibuat.",
        "Kelas 6 akan diluluskan; siswa kelas 1–5 naik 1 jenjang.",
        "Kelas tartil tidak ikut naik otomatis.",
    ])
    add_screenshot_placeholder(doc, "Halaman Tahun Ajaran")

    add_heading_custom(doc, "3.2 Semester", level=2)
    add_url(doc, f"{BASE_URL}/admin/semester")
    add_paragraph_custom(doc, "Daftar semester dalam tahun ajaran. Setiap semester memiliki status aktif/nonaktif/ditutup.")
    add_steps(doc, [
        "Klik 'Detail' pada semester.",
        "Klik 'Aktifkan' untuk mengaktifkan semester.",
        "Klik 'Refresh Snapshot' untuk memperbarui data kelas dan siswa di semester aktif.",
        "Klik 'Tutup Semester' untuk mengunci data (snapshot R2, jurnal, munaqosyah, tahfidz, riwayat kelas).",
    ])
    add_info_box(doc, "Catatan:", [
        "Tombol Refresh Snapshot hanya muncul saat semester belum ditutup.",
        "Setelah ditutup, data semester tidak dapat diubah.",
    ])
    add_screenshot_placeholder(doc, "Detail semester dengan tombol Refresh Snapshot")

    # 4. Data Siswa
    add_heading_custom(doc, "4. Manajemen Data Siswa", level=1)
    add_heading_custom(doc, "4.1 Daftar Siswa", level=2)
    add_url(doc, f"{BASE_URL}/admin/manajemen/siswa")
    add_paragraph_custom(doc, "CRUD data siswa, reset password, aktifkan/nonaktifkan, hapus, dan mutasi keluar.")
    add_steps(doc, [
        "Klik 'Data Siswa' → 'Daftar Siswa'.",
        "Klik 'Tambah Siswa' untuk input manual.",
        "Isi NIS, nama, kelas reguler, kelas tartil, nomor HP, dan status.",
    ])
    add_screenshot_placeholder(doc, "Daftar siswa")

    add_heading_custom(doc, "4.2 Import Siswa", level=2)
    add_url(doc, f"{BASE_URL}/admin/siswa/import")
    add_paragraph_custom(doc, "Impor data siswa dari file Excel sesuai template.")
    add_steps(doc, [
        "Unduh template Excel.",
        "Isi data sesuai kolom wajib.",
        "Upload file dan klik 'Import'.",
    ])
    add_screenshot_placeholder(doc, "Import siswa via Excel")

    add_heading_custom(doc, "4.3 Penempatan Tartil", level=2)
    add_url(doc, f"{BASE_URL}/admin/siswa/penempatan")
    add_paragraph_custom(doc, "Penempatan massal siswa ke kelas tartil.")
    add_screenshot_placeholder(doc, "Penempatan siswa ke kelas tartil")

    # 5. Data Guru
    add_heading_custom(doc, "5. Manajemen Data Guru", level=1)
    add_heading_custom(doc, "5.1 Guru Tartil", level=2)
    add_url(doc, f"{BASE_URL}/admin/manajemen/guru")
    add_paragraph_custom(doc, "CRUD akun guru tartil yang digunakan untuk login dan mengampu kelas.")
    add_screenshot_placeholder(doc, "Daftar guru tartil")

    add_heading_custom(doc, "5.2 Import Guru", level=2)
    add_url(doc, f"{BASE_URL}/admin/guru/import")
    add_paragraph_custom(doc, "Impor guru reguler atau tartil dari file Excel.")
    add_screenshot_placeholder(doc, "Import guru")

    add_heading_custom(doc, "5.3 Guru Reguler", level=2)
    add_url(doc, f"{BASE_URL}/admin/guru-reguler")
    add_paragraph_custom(doc, "Data guru wali kelas reguler.")
    add_screenshot_placeholder(doc, "Daftar guru reguler")

    # 6. Kelas Reguler
    add_heading_custom(doc, "6. Manajemen Kelas Reguler", level=1)
    add_heading_custom(doc, "6.1 Daftar Kelas Reguler", level=2)
    add_url(doc, f"{BASE_URL}/admin/kelas-reguler/daftar")
    add_paragraph_custom(doc, "CRUD kelas reguler (jenjang + rombel, misal: 1A, 2B).")
    add_screenshot_placeholder(doc, "Daftar kelas reguler")

    add_heading_custom(doc, "6.2 Keterangan Kelas", level=2)
    add_url(doc, f"{BASE_URL}/admin/kelas-reguler/keterangan")
    add_paragraph_custom(doc, "Lihat daftar keterangan siswa per kelas reguler.")
    add_screenshot_placeholder(doc, "Keterangan kelas reguler")

    add_heading_custom(doc, "6.3 Pindah Kelas Reguler", level=2)
    add_url(doc, f"{BASE_URL}/admin/kelas-reguler/pindah-kelas")
    add_paragraph_custom(doc, "Pindahkan siswa antar kelas reguler secara manual.")
    add_screenshot_placeholder(doc, "Pindah kelas reguler")

    # 7. Kelas Tartil
    add_heading_custom(doc, "7. Manajemen Kelas Tartil", level=1)
    add_heading_custom(doc, "7.1 Kelas Tartil", level=2)
    add_url(doc, f"{BASE_URL}/admin/kelas")
    add_paragraph_custom(doc, "CRUD kelas tartil (BQ 1, BQ 2, BQ 3, BQ 4, Tartil, Tahfidz) beserta guru pengampu.")
    add_steps(doc, [
        "Klik 'Kelas Tartil'.",
        "Klik 'Tambah Kelas'.",
        "Isi nama, mata pelajaran, jenis kelas, dan pilih guru.",
    ])
    add_screenshot_placeholder(doc, "Daftar kelas tartil")

    add_heading_custom(doc, "7.2 Pengaturan Indikator", level=2)
    add_url(doc, f"{BASE_URL}/admin/pengaturan-kelas")
    add_paragraph_custom(doc, "Atur indikator penilaian rapor internal per jenis kelas.")
    add_screenshot_placeholder(doc, "Pengaturan indikator kelas")

    add_heading_custom(doc, "7.3 Pindah Tartil (Admin View)", level=2)
    add_url(doc, f"{BASE_URL}/admin/perpindahan-tartil")
    add_paragraph_custom(doc, "Admin dapat mengajukan dan menyetujui perpindahan siswa antar kelas tartil.")
    add_screenshot_placeholder(doc, "Perpindahan kelas tartil admin")

    add_heading_custom(doc, "7.4 Rekap Kelas Tartil", level=2)
    add_url(doc, f"{BASE_URL}/admin/rekap-kelas-tartil")
    add_paragraph_custom(doc, "Rekap komposisi siswa per kelas tartil.")
    add_screenshot_placeholder(doc, "Rekap kelas tartil")

    # 8. Jurnal, Absensi, Monitoring
    add_heading_custom(doc, "8. Jurnal, Absensi, dan Monitoring", level=1)
    add_heading_custom(doc, "8.1 Progress Jurnal", level=2)
    add_url(doc, f"{BASE_URL}/admin/progress-jurnal")
    add_paragraph_custom(doc, "Monitoring pengisian jurnal harian oleh guru per tahun ajaran, semester, dan guru.")
    add_info_box(doc, "Korelasional:", [
        "Guru mengisi jurnal di /guru/jurnal.",
        "Progress jurnal admin menampilkan persentase hari efektif yang sudah terisi.",
        "Hari libur per kelas mengurangi target hari efektif.",
    ])
    add_screenshot_placeholder(doc, "Progress jurnal admin")

    add_heading_custom(doc, "8.2 Progress Absensi", level=2)
    add_url(doc, f"{BASE_URL}/admin/progress-absensi")
    add_paragraph_custom(doc, "Monitoring absensi siswa per kelas.")
    add_screenshot_placeholder(doc, "Progress absensi")

    add_heading_custom(doc, "8.3 Monitoring Guru", level=2)
    add_url(doc, f"{BASE_URL}/admin/monitoring-guru")
    add_paragraph_custom(doc, "Daftar guru yang belum mengisi jurnal, dengan hitungan hari libur per kelas.")
    add_screenshot_placeholder(doc, "Monitoring guru")

    add_heading_custom(doc, "8.4 Rekap Jurnal Bulanan", level=2)
    add_url(doc, f"{BASE_URL}/admin/jurnal-bulanan")
    add_paragraph_custom(doc, "Lihat rekap jurnal bulanan per kelas.")
    add_screenshot_placeholder(doc, "Rekap jurnal bulanan")

    add_heading_custom(doc, "8.5 Hari Libur", level=2)
    add_paragraph_custom(doc, "Admin dapat menandai hari libur per kelas melalui form yang tersedia di halaman monitoring/progress. Hari libur memengaruhi perhitungan target hari efektif.")
    add_screenshot_placeholder(doc, "Form menandai hari libur")

    # 9. Tahfidz
    add_heading_custom(doc, "9. Tahfidz & Hafalan", level=1)
    add_heading_custom(doc, "9.1 Monitoring Hafalan", level=2)
    add_url(doc, f"{BASE_URL}/admin/tahfidz")
    add_paragraph_custom(doc, "Rekap hafalan per kelas tartil. Progress ditampilkan per juz 1–30.")
    add_info_box(doc, "Aturan Hafalan:", [
        "Hafalan tidak harus berurutan (juz 30 boleh didahulukan).",
        "Status: baru, setengah_hafal, hafal, murajaah.",
        "Hafal dan murajaah dihitung sebagai ayat yang dikuasai.",
        "Hafalan bersifat kumulatif antar semester.",
    ])
    add_screenshot_placeholder(doc, "Monitoring hafalan admin")

    add_heading_custom(doc, "9.2 Rekap Hafalan per Semester", level=2)
    add_url(doc, f"{BASE_URL}/admin/tahfidz/rekap-semester")
    add_paragraph_custom(doc, "Rekap per juz per semester: total siswa, siswa sudah hafal, dan siswa tuntas (persentase ≥ 100%).")
    add_screenshot_placeholder(doc, "Rekap hafalan per semester")

    add_heading_custom(doc, "9.3 Detail Hafalan Siswa", level=2)
    add_url(doc, f"{BASE_URL}/admin/tahfidz/siswa/{{id_siswa}}")
    add_paragraph_custom(doc, "Lihat riwayat setoran, progress juz, dan surat yang sudah dihafal siswa.")
    add_screenshot_placeholder(doc, "Detail hafalan siswa")

    # 10. Munaqosyah
    add_heading_custom(doc, "10. Ujian Munaqosyah", level=1)
    add_heading_custom(doc, "10.1 Kelola Ujian", level=2)
    add_url(doc, f"{BASE_URL}/admin/munaqosyah")
    add_paragraph_custom(doc, "Admin membuat dan mengelola ujian munaqosyah.")
    add_steps(doc, [
        "Klik 'Munaqosyah' → 'Ujian'.",
        "Klik 'Tambah Ujian'.",
        "Isi nama ujian, tingkat, surat mulai/selesai, tanggal ujian, semester.",
    ])
    add_screenshot_placeholder(doc, "Daftar ujian munaqosyah")

    add_heading_custom(doc, "10.2 Approval Pendaftaran", level=2)
    add_url(doc, f"{BASE_URL}/admin/munaqosyah-approval")
    add_paragraph_custom(doc, "Setujui atau tolak pendaftaran siswa yang diajukan guru.")
    add_info_box(doc, "Alur Munaqosyah:", [
        "Admin buat ujian.",
        "Guru daftarkan siswa.",
        "Admin approve pendaftaran.",
        "Guru beri nilai lulus/tidak lulus.",
        "Siswa lihat hasil di menu Munaqosyah.",
    ])
    add_screenshot_placeholder(doc, "Approval pendaftaran munaqosyah")

    add_heading_custom(doc, "10.3 Rekap History", level=2)
    add_url(doc, f"{BASE_URL}/admin/munaqosyah-rekap")
    add_paragraph_custom(doc, "Rekap history kelulusan munaqosyah per siswa.")
    add_screenshot_placeholder(doc, "Rekap history munaqosyah")

    # 11. Penilaian Rapor
    add_heading_custom(doc, "11. Penilaian Rapor Internal", level=1)
    add_heading_custom(doc, "11.1 Buat Penilaian", level=2)
    add_url(doc, f"{BASE_URL}/admin/penilaian-rapor-internal")
    add_paragraph_custom(doc, "Admin membuat penilaian rapor internal per semester. Hanya boleh ada 1 penilaian per semester.")
    add_info_box(doc, "Korelasional:", [
        "Admin buat penilaian dan atur indikator.",
        "Guru mengisi nilai per indikator di /guru/penilaian-rapor.",
        "Sistem menghitung R2 Penilaian dan R2 Akhir.",
        "Rapor PDF dapat dicetak setelah semester ditutup.",
    ])
    add_screenshot_placeholder(doc, "Buat penilaian rapor internal")

    add_heading_custom(doc, "11.2 Progress Rapor", level=2)
    add_url(doc, f"{BASE_URL}/admin/penilaian-rapor-internal-rekap")
    add_paragraph_custom(doc, "Pantau progress pengisian nilai oleh guru.")
    add_screenshot_placeholder(doc, "Progress pengisian nilai rapor")

    add_heading_custom(doc, "11.3 Cetak Rapor", level=2)
    add_url(doc, f"{BASE_URL}/admin/cetak-rapor")
    add_paragraph_custom(doc, "Cetak PDF rapor per siswa atau per kelas.")
    add_screenshot_placeholder(doc, "Cetak rapor PDF")

    add_heading_custom(doc, "11.4 Kop Surat Rapor", level=2)
    add_url(doc, f"{BASE_URL}/admin/kop-surat-rapor")
    add_paragraph_custom(doc, "Pengaturan kop surat, logo, stempel, tanda tangan, dan catatan kaki untuk rapor PDF.")
    add_screenshot_placeholder(doc, "Pengaturan kop surat rapor")

    # 12. Perpindahan & Kenaikan
    add_heading_custom(doc, "12. Perpindahan & Kenaikan Kelas", level=1)
    add_heading_custom(doc, "12.1 Kenaikan Kelas Reguler Massal", level=2)
    add_url(doc, f"{BASE_URL}/admin/kenaikan-kelas")
    add_paragraph_custom(doc, "Proses kenaikan kelas reguler massal dan mutasi siswa keluar.")
    add_info_box(doc, "Catatan:", [
        "Kenaikan kelas reguler otomatis terjadi saat buat tahun ajaran baru.",
        "Menu ini digunakan untuk proses manual atau mutasi.",
        "Mutasi hanya dari aktif → mutasi_keluar.",
    ])
    add_screenshot_placeholder(doc, "Kenaikan kelas reguler massal")

    add_heading_custom(doc, "12.2 Riwayat Siswa", level=2)
    add_url(doc, f"{BASE_URL}/admin/riwayat-siswa")
    add_paragraph_custom(doc, "Lihat riwayat siswa per kelas reguler.")
    add_screenshot_placeholder(doc, "Riwayat siswa per kelas reguler")

    add_heading_custom(doc, "12.3 Track Record", level=2)
    add_url(doc, f"{BASE_URL}/admin/track-record")
    add_paragraph_custom(doc, "Lihat riwayat siswa per kelas tartil, termasuk perpindahan dan rekap performa per semester.")
    add_screenshot_placeholder(doc, "Track record siswa per kelas tartil")

    # 13. Pendampingan Ortu
    add_heading_custom(doc, "13. Pendampingan Orang Tua", level=1)
    add_url(doc, f"{BASE_URL}/admin/pendampingan-ortu")
    add_paragraph_custom(doc, "Monitoring laporan tadarus/murajaah yang dikirim siswa/orang tua.")
    add_info_box(doc, "Alur:", [
        "Siswa/orang tua laporkan kegiatan tadarus/murajaah di /siswa/pendampingan-ortu.",
        "Guru konfirmasi laporan di /guru/pendampingan-ortu.",
        "Admin memantau seluruh laporan di halaman ini.",
    ])
    add_screenshot_placeholder(doc, "Monitoring pendampingan orang tua")

    # 14. Audit & Statistik
    add_heading_custom(doc, "14. Audit, Statistik & Laporan", level=1)
    add_heading_custom(doc, "14.1 Audit Semester", level=2)
    add_url(doc, f"{BASE_URL}/admin/audit-semester")
    add_paragraph_custom(doc, "Lihat data yang sudah di-lock saat semester ditutup: R2, jurnal, munaqosyah, tahfidz, riwayat kelas.")
    add_screenshot_placeholder(doc, "Audit semester")

    add_heading_custom(doc, "14.2 Statistik", level=2)
    add_url(doc, f"{BASE_URL}/admin/statistik")
    add_paragraph_custom(doc, "Dashboard grafik perkembangan siswa dan kelas.")
    add_screenshot_placeholder(doc, "Dashboard statistik")

    # 15. System Setup
    add_heading_custom(doc, "15. System Setup & Maintenance", level=1)
    add_url(doc, f"{BASE_URL}/admin/system/setup")
    add_paragraph_custom(doc, "Tersembunyi dari sidebar. Digunakan untuk setup sistem, precalculate R2, clear cache, optimize, dan menjalankan artisan command tertentu.")
    add_info_box(doc, "Akses:", [
        "Buka URL langsung: /admin/system/setup.",
        "Klik 'Run Setup' untuk setup awal.",
        "Klik 'R2 Precalculate' untuk menghitung ulang cache R2.",
    ])
    add_screenshot_placeholder(doc, "System setup")

    doc.save(os.path.join(OUTPUT_DIR, "Buku_Panduan_Admin_TARTIL.docx"))


def generate_guru_doc():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Inter"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(68, 64, 60)

    create_cover(doc, "Buku Panduan Guru", "Panduan Lengkap Penggunaan Sistem untuk Guru")

    sections = [
        "Login dan Dashboard Guru",
        "Jurnal Harian & Absensi",
        "Rekap Jurnal & Absensi",
        "Tahfidz & Setoran Hafalan",
        "Ujian Munaqosyah",
        "Penilaian Rapor Internal",
        "Perpindahan Kelas Tartil",
        "Pendampingan Orang Tua",
        "Data Siswa & Profil",
    ]
    create_toc(doc, sections)

    # 1. Login
    add_heading_custom(doc, "1. Login dan Dashboard Guru", level=1)
    add_url(doc, f"{BASE_URL}/")
    add_paragraph_custom(doc, "Guru login menggunakan email dan password melalui halaman landing utama.")
    add_steps(doc, [
        "Buka URL aplikasi.",
        "Masukkan email guru dan password.",
        "Klik Masuk.",
        "Sistem mengarahkan ke dashboard guru.",
    ])
    add_screenshot_placeholder(doc, "Halaman login")

    add_heading_custom(doc, "1.1 Dashboard Guru", level=2)
    add_url(doc, f"{BASE_URL}/guru/dashboard")
    add_paragraph_custom(doc, "Menampilkan ringkasan kelas yang diampu, jadwal, dan tugas guru.")
    add_screenshot_placeholder(doc, "Dashboard guru")

    # 2. Jurnal
    add_heading_custom(doc, "2. Jurnal Harian & Absensi", level=1)
    add_url(doc, f"{BASE_URL}/guru/jurnal")
    add_paragraph_custom(doc, "Guru mengisi jurnal harian untuk setiap siswa di kelas yang diampu. Penilaian B (Baik), C (Cukup), K (Kurang).")
    add_steps(doc, [
        "Pilih tanggal jurnal.",
        "Isi informasi umum: surat, ayat, halaman, materi, topik, rencana.",
        "Beri nilai B/C/K untuk setiap siswa.",
        "Siswa yang tidak dinilai dapat di-set sakit/izin/alpha.",
        "Klik Simpan.",
    ])
    add_info_box(doc, "Korelasional:", [
        "Setelah guru menyimpan jurnal, siswa yang dinilai otomatis tercatat hadir.",
        "Siswa melihat jurnal terbaru di dashboard siswa.",
        "Admin memantau progress pengisian jurnal di /admin/progress-jurnal.",
    ])
    add_screenshot_placeholder(doc, "Halaman input jurnal harian")

    add_heading_custom(doc, "2.1 Batch Action Jurnal", level=2)
    add_paragraph_custom(doc, "Tersedia tombol aksi cepat: Copy dari kemarin, Set Hadir Semua, dan Batch Store.")
    add_screenshot_placeholder(doc, "Tombol aksi cepat jurnal")

    add_heading_custom(doc, "2.2 Jurnal Bulanan", level=2)
    add_url(doc, f"{BASE_URL}/guru/jurnal/bulanan")
    add_paragraph_custom(doc, "Lihat rekap jurnal bulanan kelas yang diampu.")
    add_screenshot_placeholder(doc, "Rekap jurnal bulanan guru")

    add_heading_custom(doc, "2.3 Rekap Absensi", level=2)
    add_url(doc, f"{BASE_URL}/guru/jurnal/rekap")
    add_paragraph_custom(doc, "Lihat rekap absensi per bulan untuk kelas yang diampu.")
    add_screenshot_placeholder(doc, "Rekap absensi guru")

    # 3. Tahfidz
    add_heading_custom(doc, "3. Tahfidz & Setoran Hafalan", level=1)
    add_url(doc, f"{BASE_URL}/guru/tahfidz")
    add_paragraph_custom(doc, "Guru mencatat setoran hafalan siswa di kelasnya. Dapat memilih juz 1–30, surat, ayat, status, dan kualitas.")
    add_steps(doc, [
        "Pilih juz yang ingin ditampilkan (opsional).",
        "Klik tombol '+ Setoran' pada siswa.",
        "Pilih juz, surat, ayat mulai/selesai.",
        "Pilih status (baru, setengah_hafal, hafal, murajaah) dan kualitas.",
        "Pilih tanggal setoran.",
        "Klik Simpan.",
    ])
    add_info_box(doc, "Korelasional:", [
        "Setoran hafalan guru akan muncul di dashboard siswa bagian Hafalan.",
        "Siswa/orang tua perlu mengkonfirmasi setoran di /siswa/hafalan.",
        "Guru dapat melihat status konfirmasi ortu di halaman ini.",
    ])
    add_info_box(doc, "Aturan Hafalan:", [
        "Siswa boleh menghafal juz tidak berurutan.",
        "Status hafal = seluruh ayat/surat sudah dikuasai.",
        "Murajaah tetap dihitung sebagai ayat yang dikuasai.",
    ])
    add_screenshot_placeholder(doc, "Halaman Tahfidz guru")

    add_heading_custom(doc, "3.1 Detail Hafalan Siswa", level=2)
    add_url(doc, f"{BASE_URL}/guru/tahfidz/siswa/{{id_siswa}}")
    add_paragraph_custom(doc, "Lihat detail progress juz dan riwayat setoran siswa.")
    add_screenshot_placeholder(doc, "Detail hafalan siswa dari sisi guru")

    # 4. Munaqosyah
    add_heading_custom(doc, "4. Ujian Munaqosyah", level=1)
    add_heading_custom(doc, "4.1 Daftar Ujian", level=2)
    add_url(doc, f"{BASE_URL}/guru/munaqosyah")
    add_paragraph_custom(doc, "Guru melihat daftar ujian munaqosyah yang sudah dibuat admin.")
    add_screenshot_placeholder(doc, "Daftar ujian munaqosyah guru")

    add_heading_custom(doc, "4.2 Pendaftaran Siswa", level=2)
    add_url(doc, f"{BASE_URL}/guru/munaqosyah/{{id_ujian}}")
    add_paragraph_custom(doc, "Guru mendaftarkan siswa kelasnya ke ujian. Pendaftaran perlu di-approve admin.")
    add_steps(doc, [
        "Klik detail ujian.",
        "Pilih siswa yang akan didaftarkan.",
        "Klik Daftar.",
        "Tunggu approval admin.",
    ])
    add_info_box(doc, "Alur Munaqosyah:", [
        "Admin buat ujian.",
        "Guru daftarkan siswa.",
        "Admin approve.",
        "Guru beri nilai lulus/tidak lulus.",
        "Siswa melihat hasil di menu Munaqosyah.",
    ])
    add_screenshot_placeholder(doc, "Pendaftaran siswa ke ujian munaqosyah")

    add_heading_custom(doc, "4.3 Approval Rekap", level=2)
    add_url(doc, f"{BASE_URL}/guru/munaqosyah/approval-rekap")
    add_paragraph_custom(doc, "Guru melihat status pendaftaran siswa: menunggu approval, disetujui, atau ditolak.")
    add_screenshot_placeholder(doc, "Rekap approval pendaftaran munaqosyah")

    # 5. Penilaian Rapor
    add_heading_custom(doc, "5. Penilaian Rapor Internal", level=1)
    add_url(doc, f"{BASE_URL}/guru/penilaian-rapor")
    add_paragraph_custom(doc, "Guru mengisi nilai rapor internal per indikator untuk siswa di kelasnya.")
    add_steps(doc, [
        "Pilih penilaian aktif.",
        "Pilih kelas.",
        "Isi nilai B/C/K per indikator untuk setiap siswa.",
        "Simpan.",
    ])
    add_info_box(doc, "Korelasional:", [
        "Admin membuat penilaian di /admin/penilaian-rapor-internal.",
        "Nilai ini menjadi komponen R2 Penilaian.",
        "R2 Akhir = (R2 Harian + R2 Penilaian) / 2.",
    ])
    add_screenshot_placeholder(doc, "Halaman isi nilai rapor")

    add_heading_custom(doc, "5.1 Rekap Nilai Rapor", level=2)
    add_url(doc, f"{BASE_URL}/guru/rekap-nilai-rapor")
    add_paragraph_custom(doc, "Lihat rekap nilai yang sudah diisi beserta R2 Harian, R2 Penilaian, dan R2 Akhir.")
    add_screenshot_placeholder(doc, "Rekap nilai rapor guru")

    # 6. Perpindahan
    add_heading_custom(doc, "6. Perpindahan Kelas Tartil", level=1)
    add_heading_custom(doc, "6.1 Pindah Individu", level=2)
    add_url(doc, f"{BASE_URL}/guru/perpindahan/create")
    add_paragraph_custom(doc, "Guru mengajukan perpindahan satu siswa ke kelas tartil lain.")
    add_steps(doc, [
        "Pilih siswa.",
        "Pilih kelas tujuan.",
        "Isi alasan.",
        "Kirim pengajuan.",
    ])
    add_screenshot_placeholder(doc, "Pengajuan perpindahan individu")

    add_heading_custom(doc, "6.2 Pindah Massal", level=2)
    add_url(doc, f"{BASE_URL}/guru/perpindahan/massal")
    add_paragraph_custom(doc, "Guru mengajukan perpindahan banyak siswa sekaligus (3 langkah).")
    add_screenshot_placeholder(doc, "Pengajuan perpindahan massal")

    add_heading_custom(doc, "6.3 Approval Pindah", level=2)
    add_url(doc, f"{BASE_URL}/guru/perpindahan/approval")
    add_paragraph_custom(doc, "Guru kelas tujuan dapat menyetujui atau menolak perpindahan masuk.")
    add_info_box(doc, "Alur Perpindahan:", [
        "Guru asal ajukan perpindahan.",
        "Admin atau guru kelas tujuan approve/tolak.",
        "Siswa pindah ke kelas tartil baru.",
        "Siswa melihat riwayat di /siswa/perpindahan.",
    ])
    add_screenshot_placeholder(doc, "Approval perpindahan kelas")

    # 7. Pendampingan Ortu
    add_heading_custom(doc, "7. Pendampingan Orang Tua", level=1)
    add_url(doc, f"{BASE_URL}/guru/pendampingan-ortu")
    add_paragraph_custom(doc, "Guru mengkonfirmasi laporan tadarus/murajaah yang dikirim siswa/orang tua.")
    add_steps(doc, [
        "Lihat daftar laporan masuk.",
        "Klik konfirmasi untuk laporan yang valid.",
        "Tambahkan catatan jika perlu.",
    ])
    add_info_box(doc, "Korelasional:", [
        "Siswa/orang tua mengirim laporan di /siswa/pendampingan-ortu.",
        "Guru konfirmasi di halaman ini.",
        "Siswa melihat status konfirmasi di halaman yang sama.",
    ])
    add_screenshot_placeholder(doc, "Konfirmasi pendampingan orang tua")

    # 8. Data Siswa & Track Record
    add_heading_custom(doc, "8. Data Siswa & Track Record", level=1)
    add_heading_custom(doc, "8.1 Data Siswa Kelas", level=2)
    add_url(doc, f"{BASE_URL}/guru/siswa")
    add_paragraph_custom(doc, "Guru melihat daftar siswa di kelas yang diampu dan dapat mengedit nomor HP siswa.")
    add_screenshot_placeholder(doc, "Data siswa kelas guru")

    add_heading_custom(doc, "8.2 Track Record Siswa", level=2)
    add_url(doc, f"{BASE_URL}/guru/track-record")
    add_paragraph_custom(doc, "Lihat riwayat kelas dan performa jurnal siswa per semester.")
    add_screenshot_placeholder(doc, "Track record siswa dari sisi guru")

    add_heading_custom(doc, "8.3 Ganti Password", level=2)
    add_url(doc, f"{BASE_URL}/guru/password/edit")
    add_paragraph_custom(doc, "Guru mengubah password akunnya.")
    add_screenshot_placeholder(doc, "Ganti password guru")

    doc.save(os.path.join(OUTPUT_DIR, "Buku_Panduan_Guru_TARTIL.docx"))


def generate_siswa_doc():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Inter"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(68, 64, 60)

    create_cover(doc, "Buku Panduan Siswa", "Panduan Lengkap Penggunaan Sistem untuk Siswa dan Orang Tua")

    sections = [
        "Login dan Navigasi Siswa",
        "Dashboard Siswa",
        "Rapor",
        "Hafalan Al-Quran",
        "Pendampingan Orang Tua",
        "Riwayat Kelas",
        "Track Record",
        "Munaqosyah",
        "Absensi",
        "Profil",
    ]
    create_toc(doc, sections)

    # 1. Login
    add_heading_custom(doc, "1. Login dan Navigasi Siswa", level=1)
    add_url(doc, f"{BASE_URL}/siswa/login")
    add_paragraph_custom(doc, "Siswa login menggunakan NIS dan nomor HP yang terdaftar di sistem.")
    add_steps(doc, [
        "Buka URL login siswa.",
        "Masukkan NIS.",
        "Masukkan nomor HP (seperti terdaftar di sistem).",
        "Klik Masuk.",
    ])
    add_info_box(doc, "Catatan:", [
        "Jika nomor HP berubah, hubungi guru atau admin untuk diperbarui.",
        "Siswa juga dapat mengedit nomor HP sendiri di menu Profil.",
    ])
    add_screenshot_placeholder(doc, "Halaman login siswa")

    add_heading_custom(doc, "1.1 Navigasi Menu", level=2)
    add_paragraph_custom(doc, "Setelah login, siswa melihat menu horizontal di bagian atas: Dashboard, Rapor, Hafalan, Pendampingan Ortu, Riwayat Kelas, Track Record, dan Profil.")
    add_screenshot_placeholder(doc, "Navigasi menu siswa")

    # 2. Dashboard
    add_heading_custom(doc, "2. Dashboard Siswa", level=1)
    add_url(doc, f"{BASE_URL}/siswa/dashboard")
    add_paragraph_custom(doc, "Dashboard menampilkan ringkasan pembelajaran siswa: R2 Harian, R2 Penilaian, R2 Akhir, progress hafalan, statistik B/C/K, grafik bulanan, dan jurnal terbaru.")
    add_info_box(doc, "Penjelasan R2:", [
        "R2 Harian = hasil penilaian jurnal harian (B=2, C=1, K=0).",
        "R2 Penilaian = hasil penilaian rapor internal.",
        "R2 Akhir = rata-rata dari R2 Harian dan R2 Penilaian.",
        "R2 Akhir Sebelum Penilaian = nilai R2 Harian.",
        "R2 Akhir Setelah Penilaian = nilai R2 Akhir gabungan.",
    ])
    add_screenshot_placeholder(doc, "Dashboard siswa")

    add_heading_custom(doc, "2.1 Jurnal Terbaru", level=2)
    add_paragraph_custom(doc, "Menampilkan 30 jurnal terbaru dengan tanggal, penilaian B/C/K, surat, dan catatan guru. Catatan guru ditampilkan langsung agar siswa/orang tua dapat memantau keterangan dari guru.")
    add_info_box(doc, "Korelasional:", [
        "Guru menginput jurnal di /guru/jurnal.",
        "Catatan guru yang ditulis saat input jurnal muncul di sini.",
    ])
    add_screenshot_placeholder(doc, "Jurnal terbaru dengan catatan guru")

    add_heading_custom(doc, "2.2 Progress Hafalan", level=2)
    add_paragraph_custom(doc, "Jika siswa memiliki kelas tartil, dashboard menampilkan progress juz 1–30 dengan warna status: hafal, murajaah, setengah hafal, baru, atau belum.")
    add_screenshot_placeholder(doc, "Progress hafalan di dashboard")

    add_heading_custom(doc, "2.3 Popup Konfirmasi Orang Tua", level=2)
    add_paragraph_custom(doc, "Jika ada setoran hafalan yang belum dikonfirmasi, popup muncul di dashboard untuk memudahkan konfirmasi.")
    add_screenshot_placeholder(doc, "Popup konfirmasi orang tua di dashboard")

    # 3. Rapor
    add_heading_custom(doc, "3. Rapor", level=1)
    add_url(doc, f"{BASE_URL}/siswa/nilai")
    add_paragraph_custom(doc, "Halaman daftar rapor semester yang sudah ditutup oleh admin.")
    add_steps(doc, [
        "Klik menu 'Rapor'.",
        "Pilih semester yang tersedia.",
        "Klik 'Lihat Rapor PDF' untuk mengunduh atau melihat rapor.",
    ])
    add_info_box(doc, "Catatan:", [
        "Rapor hanya tersedia untuk semester yang sudah ditutup.",
        "Rapor semester yang sedang berlangsung tidak dapat diunduh.",
    ])
    add_screenshot_placeholder(doc, "Halaman rapor siswa")

    add_url(doc, f"{BASE_URL}/siswa/rapor?semester_id={{id}}", label=f"{BASE_URL}/siswa/rapor?semester_id=...")
    add_paragraph_custom(doc, "URL untuk melihat atau mengunduh PDF rapor semester tertentu.")
    add_screenshot_placeholder(doc, "Preview PDF rapor")

    # 4. Hafalan
    add_heading_custom(doc, "4. Hafalan Al-Quran", level=1)
    add_url(doc, f"{BASE_URL}/siswa/hafalan")
    add_paragraph_custom(doc, "Halaman untuk memantau progress hafalan, melihat setoran, dan mengkonfirmasi setoran oleh orang tua.")
    add_heading_custom(doc, "4.1 Progress Juz", level=2)
    add_paragraph_custom(doc, "Grid juz 1–30 dengan warna sesuai status. Klik juz untuk memilih filter.")
    add_screenshot_placeholder(doc, "Progress juz 1-30")

    add_heading_custom(doc, "4.2 Surat yang Telah Dihafal", level=2)
    add_paragraph_custom(doc, "Menampilkan daftar surat yang ditandai status hafal oleh guru. Hanya surat dengan status hafal 100% yang muncul di sini.")
    add_info_box(doc, "Korelasional:", [
        "Guru mencatat setoran di /guru/tahfidz atau /admin/tahfidz.",
        "Surat dengan status hafal muncul di halaman ini.",
    ])
    add_screenshot_placeholder(doc, "Daftar surat yang telah dihafal")

    add_heading_custom(doc, "4.3 Riwayat Setoran", level=2)
    add_paragraph_custom(doc, "Tabel riwayat setoran lengkap: juz, surat, ayat, status, kualitas, tanggal setoran, tanggal konfirmasi ortu, dan nama guru.")
    add_screenshot_placeholder(doc, "Riwayat setoran hafalan")

    add_heading_custom(doc, "4.4 Konfirmasi Orang Tua", level=2)
    add_paragraph_custom(doc, "Siswa/orang tua mencentang setoran yang sudah dipantau di rumah, lalu klik konfirmasi.")
    add_steps(doc, [
        "Lihat daftar setoran yang belum dikonfirmasi.",
        "Centang setoran yang sudah dipantau.",
        "Klik 'Konfirmasi Setoran Terpilih'.",
    ])
    add_info_box(doc, "Korelasional:", [
        "Setelah dikonfirmasi, status di kolom 'Tanggal Konfirmasi Ortu' akan terisi.",
        "Guru dapat melihat status konfirmasi di /guru/tahfidz.",
    ])
    add_screenshot_placeholder(doc, "Form konfirmasi orang tua")

    # 5. Pendampingan Ortu
    add_heading_custom(doc, "5. Pendampingan Orang Tua", level=1)
    add_url(doc, f"{BASE_URL}/siswa/pendampingan-ortu")
    add_paragraph_custom(doc, "Halaman untuk melaporkan kegiatan tadarus dan murajaah yang dilakukan siswa bersama orang tua di rumah.")
    add_steps(doc, [
        "Klik menu 'Pendampingan Ortu'.",
        "Pilih jenis kegiatan: Tadarus atau Murajaah.",
        "Pilih surat.",
        "Isi ayat mulai dan ayat selesai (opsional).",
        "Pilih tanggal kegiatan.",
        "Tambahkan catatan jika perlu.",
        "Klik 'Kirim Laporan'.",
    ])
    add_info_box(doc, "Korelasional:", [
        "Laporan masuk ke guru di /guru/pendampingan-ortu.",
        "Guru konfirmasi laporan.",
        "Status laporan muncul di halaman ini setelah dikonfirmasi.",
    ])
    add_screenshot_placeholder(doc, "Form laporan pendampingan orang tua")

    # 6. Riwayat Kelas
    add_heading_custom(doc, "6. Riwayat Kelas", level=1)
    add_url(doc, f"{BASE_URL}/siswa/perpindahan")
    add_paragraph_custom(doc, "Menampilkan kelas saat ini (reguler, tartil, guru) dan riwayat perpindahan kelas tartil.")
    add_info_box(doc, "Korelasional:", [
        "Perpindahan kelas tartil diajukan guru di /guru/perpindahan/create.",
        "Disetujui admin atau guru kelas tujuan.",
        "Riwayat muncul di halaman ini.",
    ])
    add_screenshot_placeholder(doc, "Halaman riwayat kelas")

    # 7. Track Record
    add_heading_custom(doc, "7. Track Record", level=1)
    add_url(doc, f"{BASE_URL}/siswa/track-record/{{id_siswa}}")
    add_paragraph_custom(doc, "Menampilkan profil siswa, riwayat perpindahan kelas, dan rekap performa per semester (rata-rata jurnal, B/C/K, total pertemuan).")
    add_screenshot_placeholder(doc, "Halaman track record siswa")

    # 8. Munaqosyah
    add_heading_custom(doc, "8. Munaqosyah", level=1)
    add_url(doc, f"{BASE_URL}/siswa/munaqosyah")
    add_paragraph_custom(doc, "Menampilkan riwayat ujian munaqosyah yang pernah diikuti, status lulus/tidak lulus/pending, dan catatan.")
    add_info_box(doc, "Korelasional:", [
        "Admin membuat ujian di /admin/munaqosyah.",
        "Guru mendaftarkan siswa di /guru/munaqosyah.",
        "Admin menyetujui pendaftaran.",
        "Guru memberi nilai lulus/tidak lulus.",
        "Siswa melihat hasil akhir di halaman ini.",
    ])
    add_screenshot_placeholder(doc, "Halaman riwayat munaqosyah siswa")

    # 9. Absensi
    add_heading_custom(doc, "9. Absensi", level=1)
    add_url(doc, f"{BASE_URL}/siswa/absensi")
    add_paragraph_custom(doc, "Menampilkan riwayat kehadiran siswa per semester. Status: Hadir, Sakit, Izin, Alpha.")
    add_info_box(doc, "Korelasional:", [
        "Absensi dihasilkan otomatis dari jurnal harian yang diisi guru.",
        "Siswa yang dinilai di hari tersebut = Hadir.",
    ])
    add_screenshot_placeholder(doc, "Halaman absensi siswa")

    # 10. Profil
    add_heading_custom(doc, "10. Profil", level=1)
    add_url(doc, f"{BASE_URL}/siswa/no-hp/edit")
    add_paragraph_custom(doc, "Siswa dapat mengedit nomor HP sendiri.")
    add_steps(doc, [
        "Klik menu 'Profil'.",
        "Masukkan nomor HP baru.",
        "Klik 'Simpan Perubahan'.",
    ])
    add_screenshot_placeholder(doc, "Halaman edit nomor HP")

    doc.save(os.path.join(OUTPUT_DIR, "Buku_Panduan_Siswa_TARTIL.docx"))


if __name__ == "__main__":
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    generate_admin_doc()
    generate_guru_doc()
    generate_siswa_doc()
    print("Generated 3 panduan Word files in", OUTPUT_DIR)
